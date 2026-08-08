from io import BytesIO

from docx import Document as DocxDocument
from manabi_server.export.docx_export import pm_to_docx


def test_pm_to_docx_roundtrip():
    pm = {
        "type": "doc",
        "content": [
            {
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": "Cache basics"}],
            },
            {
                "type": "paragraph",
                "content": [
                    {"type": "text", "text": "A hit is "},
                    {"type": "text", "marks": [{"type": "bold"}], "text": "fast"},
                ],
            },
            {
                "type": "bulletList",
                "content": [
                    {
                        "type": "listItem",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "L1 smallest"}],
                            }
                        ],
                    }
                ],
            },
            {
                "type": "taskList",
                "content": [
                    {
                        "type": "taskItem",
                        "attrs": {"checked": True},
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [{"type": "text", "text": "review slides"}],
                            }
                        ],
                    }
                ],
            },
            {
                "type": "table",
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {
                                "type": "tableHeader",
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [{"type": "text", "text": "Level"}],
                                    }
                                ],
                            },
                            {
                                "type": "tableHeader",
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [{"type": "text", "text": "Size"}],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            },
        ],
    }
    data = pm_to_docx(pm, title="Module 4 — Notes")
    reopened = DocxDocument(BytesIO(data))
    texts = [p.text for p in reopened.paragraphs]
    assert "Module 4 — Notes" in texts
    assert "Cache basics" in texts
    assert "A hit is fast" in texts
    assert any("L1 smallest" in t for t in texts)
    assert any("☑ review slides" in t for t in texts)
    assert reopened.tables and reopened.tables[0].cell(0, 0).text == "Level"
    # the bold mark survived
    bold_runs = [r for p in reopened.paragraphs for r in p.runs if r.bold]
    assert any(r.text == "fast" for r in bold_runs)
