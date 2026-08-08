"""ProseMirror JSON → .docx (python-docx).

Covers the node set the Manabi editor produces: headings, paragraphs with
bold/italic/strike/code marks, bullet/ordered/task lists, blockquotes, code
blocks, and tables. Unknown nodes degrade to their text content.
"""

from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt

MONO = "JetBrains Mono"


def pm_to_docx(pm_json: dict, title: str) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    for node in pm_json.get("content", []) or []:
        _render_block(doc, node)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_block(doc, node: dict, list_style: str | None = None) -> None:
    ntype = node.get("type")

    if ntype == "heading":
        level = min(int(node.get("attrs", {}).get("level", 1)), 4)
        p = doc.add_heading("", level=level)
        _render_inline(p, node)
    elif ntype == "paragraph":
        p = doc.add_paragraph(style=list_style)
        _render_inline(p, node)
    elif ntype in ("bulletList", "orderedList"):
        style = "List Bullet" if ntype == "bulletList" else "List Number"
        for item in node.get("content", []) or []:
            for child in item.get("content", []) or []:
                _render_block(doc, child, list_style=style)
    elif ntype == "taskList":
        for item in node.get("content", []) or []:
            checked = item.get("attrs", {}).get("checked", False)
            box = "☑ " if checked else "☐ "
            for child in item.get("content", []) or []:
                p = doc.add_paragraph()
                p.add_run(box)
                _render_inline(p, child)
    elif ntype == "blockquote":
        for child in node.get("content", []) or []:
            p = doc.add_paragraph(style="Intense Quote")
            _render_inline(p, child)
    elif ntype == "codeBlock":
        p = doc.add_paragraph()
        run = p.add_run(_plain_text(node))
        run.font.name = MONO
        run.font.size = Pt(9)
    elif ntype == "table":
        rows = node.get("content", []) or []
        n_cols = max((len(r.get("content", []) or []) for r in rows), default=0)
        if rows and n_cols:
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row.get("content", []) or []):
                    target = table.cell(ri, ci)
                    target.text = _plain_text(cell)
    elif node.get("content"):
        for child in node["content"]:
            _render_block(doc, child, list_style=list_style)


def _render_inline(p, node: dict) -> None:
    for child in node.get("content", []) or []:
        if child.get("type") == "text":
            run = p.add_run(child.get("text", ""))
            marks = {m.get("type") for m in child.get("marks", []) or []}
            run.bold = "bold" in marks or None
            run.italic = "italic" in marks or None
            run.font.strike = "strike" in marks or None
            if "code" in marks:
                run.font.name = MONO
                run.font.size = Pt(9)
        elif child.get("type") == "hardBreak":
            p.add_run().add_break()
        else:
            _render_inline(p, child)


def _plain_text(node: dict) -> str:
    parts: list[str] = []

    def walk(n: dict) -> None:
        if n.get("type") == "text":
            parts.append(n.get("text", ""))
        for c in n.get("content", []) or []:
            walk(c)
        if n.get("type") in ("paragraph", "tableRow"):
            parts.append("\n")

    walk(node)
    return "".join(parts).strip()
